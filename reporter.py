import csv
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os
from colorama import Fore, Style

def reporter_csv(*args):
    report_csv_file = args[0]
    password_count = args[1]
    headers = ['Cracked Password', 'Count']
    # Define the target path
    directory = "Loot" 

    full_path = os.path.join(directory, report_csv_file)
    # Create the directory if it doesn't exist
    if not os.path.exists(directory):
        os.makedirs(directory)

    with open(full_path, 'w', newline='') as file:
       writer = csv.DictWriter(file, fieldnames=headers)
       writer.writeheader()
       for password, count in password_count.items():
           writer.writerow({'Cracked Password': password, 'Count': count})
    print(f"Your report is saved to: {Fore.CYAN}{full_path}{Style.RESET_ALL}!")

def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_paragraph_font(paragraph, font_name='Calibri', bold=False, size=11):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    font = run.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    return run

def reporter_docx(*args):
    report_docx_file = args[0]
    password_count = args[1]

    # Define the target path
    directory = "Loot" 

    full_path = os.path.join(directory, report_docx_file)
    # Create the directory if it doesn't exist
    if not os.path.exists(directory):
        os.makedirs(directory)

    if not isinstance(password_count, dict):
        raise ValueError("password_count must be a dictionary")

    doc = Document()
    table = doc.add_table(rows=2 + len(password_count), cols=2)
    table.style = 'Table Grid'

    # Header row
    top_row = table.rows[0]
    top_row.cells[0].merge(top_row.cells[1])
    hdr_cell = top_row.cells[0]
    hdr_cell.text = input("Enter the top-level header for the report table: ")
    paragraph = hdr_cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(paragraph, font_name='Calibri', bold=False, size=12)
    set_cell_background(hdr_cell, 'BFBFBF')

    # Column headers
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = "Password Cracked"
    hdr_cells[1].text = "Count"
    for cell in hdr_cells:
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_font(paragraph, font_name='Calibri', bold=False)
    # Data rows
    for i, (col1, col2) in enumerate(password_count.items(), start=2):
        row_cells = table.rows[i].cells
        row_cells[0].text = col1
        row_cells[1].text = str(col2)

        for cell in row_cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_font(paragraph, font_name='Calibri')

    doc.save(full_path)
    print(f"Your report is saved to: {Fore.CYAN}.\{full_path}{Style.RESET_ALL}!")
